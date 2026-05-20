# -*- coding: utf-8 -*-
"""
srs_format.py — SRS document format standard (LEXcentra)
========================================================

Decoded directly from `FINAL_SRS_IEEE_LEX.docx` — specifies the ENTIRE
presentation as Python code (`python-docx`).

PRINCIPLE: content is written in .md  ·  format/style is written in Python (this file).

Features:
  - Cover page (logo + title + metadata table)
  - Native table of contents (TOC field, auto-updates when opened in Word)
  - Heading 1..6 + Heading 1 underline
  - Automatic multi-level heading numbering (1, 1.1, 1.1.1, ...)
  - Standard SRS tables · callout · Loại/Thuộc tính legend
  - Footer with PAGE / NUMPAGES fields

Install:  pip install python-docx

OOXML units:  1 pt = 20 dxa (twips) · 1 inch = 1440 dxa · border sz = 1/8 pt
"""

import os
from docx import Document
from docx.shared import Pt, Mm, RGBColor, Twips, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# 1. STANDARD COLOR PALETTE  (styles.xml + theme1.xml)
# ============================================================
class Color:
    PRIMARY        = RGBColor(0x19, 0x3D, 0x74)  # #193D74 - H1/H2 + table header fill + H1 underline
    ACCENT_TEAL    = RGBColor(0x15, 0x60, 0x82)  # #156082 - Heading 5
    HEADING4_GRAY  = RGBColor(0x65, 0x66, 0x68)  # #656668 - Heading 4 + cover version line
    HEADING6_BLUE  = RGBColor(0x1F, 0x4D, 0x78)  # #1F4D78 - Heading 6
    CAPTION_NAVY   = RGBColor(0x0E, 0x28, 0x41)  # #0E2841 - Caption (Hình ...)
    BODY_TEXT      = RGBColor(0x25, 0x27, 0x29)  # #252729 - default text color
    TABLE_TEXT     = RGBColor(0x00, 0x00, 0x00)  # #000000 - table cell text
    HEADER_TEXT    = RGBColor(0xFF, 0xFF, 0xFF)  # #FFFFFF - table header text
    CALLOUT_FILL   = "FFF8DF"                    # #FFF8DF - callout fill "[Đang cập nhật]"
    TABLE_HDR_FILL = "193D74"                    # table header fill (hex)
    H1_BORDER      = "193D74"                    # Heading 1 underline color


# ============================================================
# 2. STANDARD FONTS
# ============================================================
class Font:
    DEFAULT = "Mulish"
    SIZE_BODY     = 11.0   # body text
    SIZE_TABLE    = 10.0   # table text
    SIZE_TITLE    = 28.0   # cover title
    SIZE_SUBTITLE = 18.0   # cover subtitle
    SIZE_VERSION  = 12.0   # cover version line
    SIZE_TOC_HEAD = 16.0   # the "Mục lục" heading text
    SIZE_H1, SIZE_H2, SIZE_H3 = 16.0, 13.0, 11.5
    SIZE_H4, SIZE_H5, SIZE_H6 = 11.0, 11.0, 11.0
    SIZE_CAPTION  = 9.0


# ============================================================
# 3. PAGE SETUP  (sectPr)
# ============================================================
class Page:
    WIDTH_MM, HEIGHT_MM = 210.0, 297.0          # A4 portrait
    MARGIN_MM = 19.05                            # 1080 dxa, all four sides
    HEADER_DIST_MM = FOOTER_DIST_MM = 12.5       # 708 dxa


# ============================================================
# 4. HEADING  (id, name, size, bold, color, before_pt, after_pt)
# ============================================================
HEADINGS = [
    ("Title",    "Title",     Font.SIZE_TITLE, False, None,                 0,  0),
    ("Heading1", "Heading 1", Font.SIZE_H1,    True,  Color.PRIMARY,        18, 8),
    ("Heading2", "Heading 2", Font.SIZE_H2,    True,  Color.PRIMARY,        14, 6),
    ("Heading3", "Heading 3", Font.SIZE_H3,    True,  None,                 10, 4),
    ("Heading4", "Heading 4", Font.SIZE_H4,    True,  Color.HEADING4_GRAY,  14, 9),
    ("Heading5", "Heading 5", Font.SIZE_H5,    True,  Color.ACCENT_TEAL,    12, 12),
    ("Heading6", "Heading 6", Font.SIZE_H6,    False, Color.HEADING6_BLUE,  0,  0),
    ("Caption",  "Caption",   Font.SIZE_CAPTION, False, Color.CAPTION_NAVY, 0,  10),
]

# Heading 1 underline (decoded: <w:pBdr><w:bottom .../></w:pBdr>)
H1_BORDER = dict(val="single", color=Color.H1_BORDER, sz="6", space="4")


# ============================================================
# 5. TABLE  (Sli Table Style)
# ============================================================
class Table:
    CONTENT_WIDTH_MM = 171.9     # 9746 dxa = A4 - 2*margin
    BORDER_SIZE_8THPT = 4        # 0.5pt
    BORDER_COLOR = "auto"
    HEADER_FILL = Color.TABLE_HDR_FILL
    HEADER_BOLD = True
    HEADER_TEXT_COLOR = Color.HEADER_TEXT


# ============================================================
# 6. MULTI-LEVEL HEADING NUMBERING
#    dedicated numId, format: 1 / 1.1 / 1.1.1 ... up to 6 levels
# ============================================================
HEADING_NUM_ID = 77   # numId dedicated to the heading multilevel list


# ============================================================
# 7. FOOTER
# ============================================================
class Footer:
    # {project} = project code/name, passed via new_srs_document(project_name=...)
    LEFT_TEXT_TEMPLATE = "{project}  ·  Software Requirements Specification"
    DEFAULT_PROJECT = "[Tên dự án]"
    PAGE_PREFIX = "Trang "
    PAGE_SEP = " / "


# ============================================================
# 8. COVER PAGE
# ============================================================
class Cover:
    # logo extracted from the original SRS (image1.png), at its native display size
    LOGO_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                             "assets", "srs_logo.png")
    LOGO_WIDTH_MM = 90.0
    LOGO_HEIGHT_MM = 19.6
    SPACE_AFTER_LOGO_PT = 24
    SPACE_AFTER_TITLE_PT = 4
    SPACE_AFTER_SUBTITLE_PT = 2
    SPACE_AFTER_VERSION_PT = 18


# ============================================================
# 9. LEGEND — component Loại & Thuộc tính reference
#    CANONICAL data source: reuse it, do NOT create new types/attributes.
# ============================================================
LEGEND_TYPES = [
    ("Input (Text/Email/Password/Search)", "Ô nhập liệu một dòng, phân loại theo kiểu dữ liệu."),
    ("Textarea", "Ô nhập liệu nhiều dòng cho nội dung dài."),
    ("Select (Single / Multi)", "Dropdown chọn một / nhiều giá trị từ danh sách."),
    ("Select + Text Input", "Dropdown có gợi ý, cho phép nhập giá trị mới nếu chưa tồn tại."),
    ("Button", "Nút hành động. Biến thể: Primary (chính), Danger (xóa/cảnh báo), Icon Button."),
    ("Checkbox / Checklist", "Ô tích chọn bật/tắt một hoặc nhiều tùy chọn."),
    ("Date Picker / Time Picker", "Bộ chọn ngày / giờ."),
    ("Toggle Button", "Nút chuyển đổi nhanh giữa các chế độ hiển thị."),
    ("Status Tag / Priority Tag / Tag", "Nhãn hiển thị trạng thái, độ ưu tiên, hoặc phân loại."),
    ("Avatar", "Ảnh đại diện người dùng, có thể kèm tên."),
    ("Text", "Văn bản hiển thị. Biến thể: Static (cố định), Inline Edit, Text + Icon, Text Color."),
    ("Label", "Nhãn mô tả ngắn cho trạng thái hoặc nhóm thông tin."),
]

LEGEND_ATTRS = [
    ("Required", "Trường bắt buộc nhập; chặn submit nếu để trống."),
    ("Required (Conditional)", "Bắt buộc khi thỏa một điều kiện cụ thể (vd đã thêm dòng)."),
    ("Unique", "Giá trị phải duy nhất trên toàn hệ thống."),
    ("Read-only", "Chỉ hiển thị, người dùng không chỉnh sửa được."),
    ("Max_N_char", "Giới hạn tối đa N ký tự cho trường."),
    ("[Mặc định]", "Đánh dấu giá trị / tiêu chí được chọn mặc định."),
]


# ============================================================
# ===================  STYLE-BUILDING FUNCTIONS  =============
# ============================================================
def _set_spacing(style, before_pt, after_pt):
    pf = style.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)


def _apply_default(doc):
    """docDefaults: Mulish 11pt #252729."""
    normal = doc.styles["Normal"]
    normal.font.name = Font.DEFAULT
    normal.font.size = Pt(Font.SIZE_BODY)
    normal.font.color.rgb = Color.BODY_TEXT
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), Font.DEFAULT)


def _add_heading1_border(style):
    """Heading 1 underline: <w:pBdr><w:bottom .../></w:pBdr>."""
    ppr = style.element.get_or_add_pPr()
    # remove any existing pBdr
    for old in ppr.findall(qn("w:pBdr")):
        ppr.remove(old)
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), H1_BORDER["val"])
    bottom.set(qn("w:sz"), H1_BORDER["sz"])
    bottom.set(qn("w:space"), H1_BORDER["space"])
    bottom.set(qn("w:color"), H1_BORDER["color"])
    pbdr.append(bottom)
    ppr.append(pbdr)


def _apply_headings(doc):
    """Load Title + Heading 1..6 + Caption; underline Heading 1."""
    for sid, name, size, bold, color, before, after in HEADINGS:
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        f = style.font
        f.name = Font.DEFAULT
        f.size = Pt(size)
        f.bold = bold
        # color None (Title, Heading 3) -> force to body color to override python-docx's default blue
        f.color.rgb = color if color is not None else Color.BODY_TEXT
        if name == "Caption":
            f.italic = True
        _set_spacing(style, before, after)
        if name == "Heading 1":
            _add_heading1_border(style)


def _apply_heading_numbering(doc):
    """
    Automatic multi-level numbering for Heading 1..6: 1 / 1.1 / 1.1.1 ...
    Creates abstractNum + num in numbering.xml, links numPr into each Heading style.
    """
    numbering = doc.part.numbering_part.element
    nsW = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    # pick an unused abstractNumId
    used_abs = [int(a.get(qn("w:abstractNumId")))
                for a in numbering.findall(qn("w:abstractNum"))]
    abs_id = (max(used_abs) + 1) if used_abs else 0

    abstractNum = OxmlElement("w:abstractNum")
    abstractNum.set(qn("w:abstractNumId"), str(abs_id))
    mlt = OxmlElement("w:multiLevelType")
    mlt.set(qn("w:val"), "multilevel")
    abstractNum.append(mlt)

    # Levels 1-4 (ilvl 0-3): decimal             ->  1 / 1.1 / 1.1.1 / 1.1.1.1
    # Level 5  (ilvl 4)    : uppercase letters   ->  A. B. C.  (restarts per parent H4)
    # Level 6  (ilvl 5)    : lowercase letters   ->  a. b. c.  (restarts per parent H5)
    for ilvl in range(6):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))
        start = OxmlElement("w:start"); start.set(qn("w:val"), "1")
        numFmt = OxmlElement("w:numFmt")
        lvlText = OxmlElement("w:lvlText")
        if ilvl <= 3:
            numFmt.set(qn("w:val"), "decimal")
            lvlText.set(qn("w:val"), ".".join(f"%{k+1}" for k in range(ilvl + 1)))
        elif ilvl == 4:
            numFmt.set(qn("w:val"), "upperLetter")
            lvlText.set(qn("w:val"), "%5.")
        else:  # ilvl == 5
            numFmt.set(qn("w:val"), "lowerLetter")
            lvlText.set(qn("w:val"), "%6.")
        lvlJc = OxmlElement("w:lvlJc"); lvlJc.set(qn("w:val"), "left")
        suff = OxmlElement("w:suff"); suff.set(qn("w:val"), "space")
        pPr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "0"); ind.set(qn("w:hanging"), "0")
        pPr.append(ind)
        for el in (start, numFmt, suff, lvlText, lvlJc, pPr):
            lvl.append(el)
        abstractNum.append(lvl)

    # insert abstractNum BEFORE the first <w:num> (correct schema order)
    first_num = numbering.find(qn("w:num"))
    if first_num is not None:
        first_num.addprevious(abstractNum)
    else:
        numbering.append(abstractNum)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(HEADING_NUM_ID))
    absRef = OxmlElement("w:abstractNumId")
    absRef.set(qn("w:val"), str(abs_id))
    num.append(absRef)
    numbering.append(num)

    # link numPr into each Heading style
    for ilvl, name in enumerate(["Heading 1", "Heading 2", "Heading 3",
                                 "Heading 4", "Heading 5", "Heading 6"]):
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        ppr = style.element.get_or_add_pPr()
        for old in ppr.findall(qn("w:numPr")):
            ppr.remove(old)
        numPr = OxmlElement("w:numPr")
        ilvlEl = OxmlElement("w:ilvl"); ilvlEl.set(qn("w:val"), str(ilvl))
        numIdEl = OxmlElement("w:numId"); numIdEl.set(qn("w:val"), str(HEADING_NUM_ID))
        numPr.append(ilvlEl); numPr.append(numIdEl)
        ppr.insert(0, numPr)


def suppress_numbering(paragraph):
    """Disable numbering for one specific paragraph (numId=0) — used for the frontmatter heading."""
    ppr = paragraph._p.get_or_add_pPr()
    for old in ppr.findall(qn("w:numPr")):
        ppr.remove(old)
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId"); numId.set(qn("w:val"), "0")
    numPr.append(ilvl); numPr.append(numId)
    ppr.insert(0, numPr)


def _apply_page(doc):
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(Page.WIDTH_MM)
        section.page_height = Mm(Page.HEIGHT_MM)
        section.top_margin = section.bottom_margin = Mm(Page.MARGIN_MM)
        section.left_margin = section.right_margin = Mm(Page.MARGIN_MM)
        section.header_distance = Mm(Page.HEADER_DIST_MM)
        section.footer_distance = Mm(Page.FOOTER_DIST_MM)


def _add_field(paragraph, field_code, placeholder="", size=None):
    """Insert a field (PAGE, NUMPAGES, TOC...) into the paragraph."""
    run = paragraph.add_run()
    if size is not None:
        run.font.size = size
    fc1 = OxmlElement("w:fldChar"); fc1.set(qn("w:fldCharType"), "begin")
    if field_code.strip().startswith("TOC"):
        fc1.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fc2 = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "separate")
    fc3 = OxmlElement("w:fldChar"); fc3.set(qn("w:fldCharType"), "end")
    run._r.append(fc1); run._r.append(instr); run._r.append(fc2)
    if placeholder:
        ph = paragraph.add_run(placeholder)
        ph._r.append(OxmlElement("w:noProof"))
    run3 = paragraph.add_run()
    if size is not None:
        run3.font.size = size
    run3._r.append(fc3)


def _apply_footer(doc, project_name):
    left_text = Footer.LEFT_TEXT_TEMPLATE.format(project=project_name)
    sz = Pt(Font.SIZE_CAPTION)
    # Remove the default tab of the "Footer" style (it has a built-in center-tab
    # that makes "\t" jump to the page center instead of the right edge).
    try:
        fppr = doc.styles["Footer"].element.find(qn("w:pPr"))
        if fppr is not None:
            for t in fppr.findall(qn("w:tabs")):
                fppr.remove(t)
    except KeyError:
        pass
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = ""
        tabs = p.paragraph_format.tab_stops
        tabs.add_tab_stop(Twips(int(Mm(Table.CONTENT_WIDTH_MM).twips)),
                          WD_TAB_ALIGNMENT.RIGHT)
        r1 = p.add_run(left_text); r1.font.size = sz
        r2 = p.add_run("\t" + Footer.PAGE_PREFIX); r2.font.size = sz
        _add_field(p, "PAGE", size=sz)
        r3 = p.add_run(Footer.PAGE_SEP); r3.font.size = sz
        _add_field(p, "NUMPAGES", size=sz)


def _enable_update_fields(doc):
    """Set the flag so Word auto-updates fields (TOC) when the file is opened."""
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        uf = OxmlElement("w:updateFields")
        uf.set(qn("w:val"), "true")
        settings.append(uf)


# ============================================================
# ===================  MAIN API  =============================
# ============================================================
def new_srs_document(project_name=None):
    """A Document preloaded with the full SRS format standard (page, font,
    heading, numbering, footer, auto-update fields).

    project_name: project code/name shown in the footer. If None -> placeholder.
    """
    doc = Document()
    _apply_default(doc)
    _apply_headings(doc)
    _apply_heading_numbering(doc)
    _apply_page(doc)
    _apply_footer(doc, project_name or Footer.DEFAULT_PROJECT)
    _enable_update_fields(doc)
    return doc


def _runs(paragraph, text, size_pt=None, color=None, bold=None, italic=None):
    run = paragraph.add_run(text)
    run.font.name = Font.DEFAULT
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def add_cover_header(doc, title, subtitle, version_line, logo_path=None):
    """
    Build the top of the cover page: logo (centered) + title + subtitle + version.
    The metadata table is rendered separately by the caller (to handle **bold** in cells).
    """
    logo_path = logo_path or Cover.LOGO_PATH

    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_after = Pt(Cover.SPACE_AFTER_LOGO_PT)
    if os.path.exists(logo_path):
        run = p_logo.add_run()
        run.add_picture(logo_path, width=Mm(Cover.LOGO_WIDTH_MM),
                        height=Mm(Cover.LOGO_HEIGHT_MM))

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(Cover.SPACE_AFTER_TITLE_PT)
    _runs(p_title, title, size_pt=Font.SIZE_TITLE, color=Color.PRIMARY, bold=True)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(Cover.SPACE_AFTER_SUBTITLE_PT)
    _runs(p_sub, subtitle, size_pt=Font.SIZE_SUBTITLE, bold=True)

    p_ver = doc.add_paragraph()
    p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ver.paragraph_format.space_after = Pt(Cover.SPACE_AFTER_VERSION_PT)
    _runs(p_ver, version_line, size_pt=Font.SIZE_VERSION,
          color=Color.HEADING4_GRAY)


def add_toc(doc):
    """Insert 'Mục lục' + a native TOC field (levels 1-3, auto-updates when opened in Word)."""
    p_head = doc.add_paragraph()
    p_head.paragraph_format.space_before = Pt(6)
    p_head.paragraph_format.space_after = Pt(6)
    _runs(p_head, "Mục lục", size_pt=Font.SIZE_TOC_HEAD,
          color=Color.PRIMARY, bold=True)
    p_toc = doc.add_paragraph()
    _add_field(p_toc, 'TOC \\o "1-3" \\h \\z \\u',
               placeholder="Nhấn chuột phải vào đây và chọn Update Field để hiện Mục lục.")


def add_page_break(doc):
    """Break to a new page."""
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _set_cell_bg(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(Table.BORDER_SIZE_8THPT))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), Table.BORDER_COLOR)
        borders.append(el)
    tblPr.append(borders)


def add_srs_table(doc, headers, rows, col_widths_mm=None):
    """Standard SRS table: 0.5pt borders, header fill #193D74 with white bold 10pt text, body 10pt."""
    n = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)

    for j, htext in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        _set_cell_bg(cell, Table.HEADER_FILL)
        _runs(cell.paragraphs[0], str(htext), size_pt=Font.SIZE_TABLE,
              color=Table.HEADER_TEXT_COLOR, bold=True)

    for i, row in enumerate(rows, start=1):
        for j in range(n):
            cell = table.rows[i].cells[j]
            cell.text = ""
            val = row[j] if j < len(row) else ""
            _runs(cell.paragraphs[0], "" if val is None else str(val),
                  size_pt=Font.SIZE_TABLE, color=Color.TABLE_TEXT)

    if col_widths_mm:
        for j, w in enumerate(col_widths_mm):
            for r in table.rows:
                r.cells[j].width = Mm(w)
    return table


def add_callout(doc, text):
    """Single-cell callout table '[Đang cập nhật]' — cream fill #FFF8DF."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)
    cell = table.rows[0].cells[0]
    cell.text = ""
    _set_cell_bg(cell, Color.CALLOUT_FILL)
    _runs(cell.paragraphs[0], text, size_pt=Font.SIZE_TABLE,
          color=Color.TABLE_TEXT, bold=True)
    return table


def add_legend_section(doc):
    """
    Insert the 2 reference tables: component Loại & Thuộc tính.
    Used ONCE in the 'Yêu cầu Hành vi Chung' section. The main document
    only REUSES these values; it never defines new ones.
    The 2 table labels use bold paragraphs (NOT Headings) -> they are not
    numbered and do not leak into the table of contents.
    """
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(6)
    _runs(p1, "Loại thành phần", size_pt=Font.SIZE_BODY,
          color=Color.PRIMARY, bold=True)
    add_srs_table(doc, headers=["Loại", "Định nghĩa"], rows=LEGEND_TYPES)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(6)
    _runs(p2, "Thuộc tính", size_pt=Font.SIZE_BODY,
          color=Color.PRIMARY, bold=True)
    add_srs_table(doc, headers=["Thuộc tính", "Định nghĩa"], rows=LEGEND_ATTRS)


# ============================================================
# SMOKE TEST
# ============================================================
if __name__ == "__main__":
    doc = new_srs_document()
    add_cover_header(
        doc,
        title="RoomBooking",
        subtitle="Software Requirements Specification",
        version_line="Phiên bản Draft 1.0.0 · 19/05/2026",
    )
    add_srs_table(doc, headers=["Dự án", ""],
                  rows=[["Loại tài liệu", "SRS"], ["Ngày", "19/05/2026"]])
    add_toc(doc)
    add_page_break(doc)
    h = doc.add_paragraph("Lịch sử Phiên bản", style="Heading 1")
    suppress_numbering(h)
    doc.add_paragraph("Giới thiệu", style="Heading 1")
    doc.add_paragraph("Mô tả Tổng quan", style="Heading 1")
    doc.add_paragraph("Bối cảnh", style="Heading 2")
    add_legend_section(doc)
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       "_srs_format_smoketest.docx")
    doc.save(out)
    print(f"Smoke test OK -> {out}")

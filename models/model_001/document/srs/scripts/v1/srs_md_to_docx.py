# -*- coding: utf-8 -*-
"""
srs_md_to_docx.py — Sinh file .docx SRS từ Markdown nội dung + chuẩn srs_format.

Dùng:
    python srs_md_to_docx.py <input.md> <output.docx>
    (mặc định: SRS_Sample.md -> SRS_Sample.docx)

Xử lý:
  - Frontmatter -> trang bìa (logo + tiêu đề) + Mục lục (TOC native) + ngắt trang
  - Heading 1 đầu tiên (Lịch sử Phiên bản) -> không numbering (frontmatter)
  - #..#####       -> Heading 1..5 (numbering tự động 1, 1.1, 1.1.1...)
  - | ... |        -> bảng SRS; hàng có các ô GIỐNG NHAU -> merge thành 1 ô
  - STT/ID rỗng của bảng Đặc tả thành phần -> COM-<heading>-<NNN>
  - Mã BR rỗng     -> BR-<heading>-<NNN>
  - "Hình ..."     -> caption căn giữa, đánh số Hình <heading>-<n>
  - {{LEGEND}}     -> chèn bảng chú giải Loại & Thuộc tính
  - **text**       -> in đậm · <br> -> xuống dòng trong ô
"""
import re
import sys
import os

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)

from srs_format import (new_srs_document, add_cover_header, add_toc,
                        add_page_break, add_legend_section, add_callout,
                        suppress_numbering, Color, Font, Table,
                        _set_table_borders, _set_cell_bg, _runs)
from docx.shared import Pt, Mm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

HEADING_STYLE = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3",
                 4: "Heading 4", 5: "Heading 5", 6: "Heading 6"}


# ---------- parse markdown ----------
def parse_md(path):
    with open(path, encoding="utf-8") as f:
        lines = f.read().splitlines()
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.strip()
        if not s:
            i += 1
            continue
        m = re.match(r"^(#{1,6})\s+(.+)$", s)
        if m:
            blocks.append(("heading", len(m.group(1)), m.group(2).strip()))
            i += 1
            continue
        if s.startswith("|"):
            tbl = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl.append(lines[i].strip())
                i += 1
            rows = []
            for r in tbl:
                cells = [c.strip() for c in r.strip("|").split("|")]
                if all(set(c) <= set("-: ") and c for c in cells):
                    continue
                rows.append(cells)
            if rows:
                blocks.append(("table", rows[0], rows[1:]))
            continue
        m = re.match(r"^(\s*)-\s+(.+)$", line)
        if m:
            blocks.append(("bullet", len(m.group(1)) // 2, m.group(2).strip()))
            i += 1
            continue
        blocks.append(("para", 0, s))
        i += 1
    return blocks


def strip_md(t):
    return t.replace("**", "").strip()


# ---------- inline **bold** + <br> ----------
def add_runs(paragraph, text, size_pt=None, color=None, force_bold=False):
    for si, seg in enumerate(text.split("<br>")):
        if si > 0:
            paragraph.add_run().add_break()
        seg = seg.replace("\\|", "|")
        for part in re.split(r"(\*\*[^*]+\*\*)", seg):
            if not part:
                continue
            bold = force_bold
            txt = part
            if part.startswith("**") and part.endswith("**"):
                bold = True
                txt = part[2:-2]
            run = paragraph.add_run(txt)
            run.font.name = Font.DEFAULT
            run.bold = bold
            if size_pt:
                run.font.size = Pt(size_pt)
            if color is not None:
                run.font.color.rgb = color


# ---------- table rendering ----------
def _row_all_same(cells):
    """Row có mọi ô (khác rỗng) giống hệt nhau -> cần merge."""
    vals = [c.strip() for c in cells if c and c.strip()]
    return len(vals) == len(cells) and len(set(vals)) == 1 and len(cells) > 1


def _merge_row(table, ri):
    """Merge toàn bộ ô của hàng ri thành 1 ô; trả về ô đã merge."""
    cells = table.rows[ri].cells
    merged = cells[0]
    for c in cells[1:]:
        merged = merged.merge(c)
    return merged


def render_table(doc, headers, rows, id_prefix=None, id_heading="", id_start=0):
    """Render bảng SRS. Nếu id_prefix set: tự sinh ID cột đầu, đánh số tiếp
    từ id_start. Trả về số thứ tự ID cuối cùng đã dùng (để caller nối tiếp)."""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)

    widths = ([Table.CONTENT_WIDTH_MM * 0.30, Table.CONTENT_WIDTH_MM * 0.70]
              if ncols == 2 else [Table.CONTENT_WIDTH_MM / ncols] * ncols)

    # ----- header row -----
    if _row_all_same(headers):
        cell = _merge_row(table, 0)
        cell.text = ""
        _set_cell_bg(cell, Table.HEADER_FILL)
        add_runs(cell.paragraphs[0], headers[0], size_pt=Font.SIZE_TABLE,
                 color=Color.HEADER_TEXT, force_bold=True)
    else:
        for j, htext in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = ""
            _set_cell_bg(cell, Table.HEADER_FILL)
            add_runs(cell.paragraphs[0], htext, size_pt=Font.SIZE_TABLE,
                     color=Color.HEADER_TEXT, force_bold=True)

    # ----- body rows -----
    id_seq = id_start
    for ri, row in enumerate(rows, start=1):
        # group row (mọi ô giống nhau) -> merge
        if _row_all_same(row):
            cell = _merge_row(table, ri)
            cell.text = ""
            add_runs(cell.paragraphs[0], row[0], size_pt=Font.SIZE_TABLE,
                     color=Color.TABLE_TEXT, force_bold=True)
            continue
        # auto STT/ID
        row = list(row) + [""] * (ncols - len(row))
        if id_prefix:
            id_seq += 1
            row[0] = f"{id_prefix}-{id_heading}-{id_seq:03d}"
        for j in range(ncols):
            cell = table.rows[ri].cells[j]
            cell.text = ""
            add_runs(cell.paragraphs[0], row[j] if row[j] else "",
                     size_pt=Font.SIZE_TABLE, color=Color.TABLE_TEXT)

    for j, w in enumerate(widths):
        for r in table.rows:
            r.cells[j].width = Mm(w)
    return id_seq


def _id_column_prefix(header0, component_area):
    """Xác định prefix ID dựa trên ô header đầu tiên."""
    h = header0.strip().lower().replace("*", "")
    if component_area and ("stt" in h or h in ("id", "stt / id", "stt/id")):
        return "COM"
    if h == "mã br":
        return "BR"
    return None


# ---------- build ----------
def build(md_path, docx_path):
    blocks = parse_md(md_path)

    # ===== frontmatter — parse TRƯỚC để lấy tên dự án cho footer =====
    idx = 0
    fm_paras, fm_table = [], None
    while idx < len(blocks) and blocks[idx][0] != "heading":
        kind, a, b = blocks[idx]
        if kind == "para":
            fm_paras.append(a if isinstance(a, str) else b)
        elif kind == "table":
            fm_table = (a, b)
        idx += 1

    title = strip_md(fm_paras[0]) if len(fm_paras) > 0 else "Untitled"
    subtitle = strip_md(fm_paras[1]) if len(fm_paras) > 1 else ""
    version = strip_md(fm_paras[2]) if len(fm_paras) > 2 else ""

    # Tên dự án = tiêu đề trang bìa -> truyền vào footer (không hard-code)
    doc = new_srs_document(project_name=title)

    add_cover_header(doc, title, subtitle, version)
    if fm_table:
        render_table(doc, fm_table[0], fm_table[1])
    add_page_break(doc)          # trang bìa đứng độc lập
    add_toc(doc)
    add_page_break(doc)          # Mục lục tách khỏi "Lịch sử Phiên bản"

    # ===== nội dung chính =====
    counter = [0] * 6
    last_h4_number = ""          # số của Heading 4 gần nhất (cho ID & Hình)
    first_heading_done = False
    pending_history_break = False  # cờ: chèn page break trước heading sau Lịch sử PB
    component_area = False
    figure_area = False           # True trong mục Wireframe HOẶC Sơ đồ luồng
    fig_seq = 0
    id_counters = {}             # (prefix, h4) -> số thứ tự ID cuối cùng

    for kind, a, b in blocks[idx:]:

        if kind == "heading":
            level, text = a, b
            # heading đầu tiên = frontmatter "Lịch sử Phiên bản" -> không numbering
            if not first_heading_done:
                first_heading_done = True
                pending_history_break = True
                p = doc.add_paragraph(text, style="Heading 1")
                suppress_numbering(p)
                component_area = figure_area = False
                continue
            # ngắt trang để "Lịch sử Phiên bản" đứng độc lập (trước "Giới thiệu")
            if pending_history_break:
                add_page_break(doc)
                pending_history_break = False
            # heading có numbering
            counter[level - 1] += 1
            for k in range(level, 6):
                counter[k] = 0
            if level == 4:
                last_h4_number = ".".join(str(counter[k]) for k in range(4))
            doc.add_paragraph(text, style=HEADING_STYLE.get(level, "Heading 6"))
            # cập nhật vùng (chỉ heading cấp <=4 mới đổi vùng)
            if level <= 4:
                low = text.lower()
                component_area = text.strip().startswith("Đặc tả các thành phần")
                figure_area = ("wireframe" in low) or ("sơ đồ luồng" in low)
                fig_seq = 0
            continue

        if kind == "table":
            headers, rows = a, b
            # bảng 1 cột -> callout (nền kem) — vd placeholder Sơ đồ luồng
            if len(headers) == 1:
                add_callout(doc, strip_md(headers[0]))
                continue
            prefix = _id_column_prefix(headers[0], component_area)
            if prefix:
                # mã code lấy theo Heading 4; đánh số nối tiếp trong cùng H4
                h4 = last_h4_number.replace(".", "")
                key = (prefix, h4)
                start = id_counters.get(key, 0)
                last = render_table(doc, headers, rows, id_prefix=prefix,
                                    id_heading=h4, id_start=start)
                id_counters[key] = last
            else:
                render_table(doc, headers, rows)
            continue

        if kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, b)
            continue

        if kind == "para":
            text = a if isinstance(a, str) else b
            # marker chèn legend
            if text.strip() == "{{LEGEND}}":
                add_legend_section(doc)
                continue
            # caption Hình trong vùng wireframe -> căn giữa + đánh số lại
            if figure_area and re.match(r"^Hình\b", text.strip()):
                fig_seq += 1
                desc = re.sub(r"^Hình\s*[\w.\-]*\.?\s*", "", text.strip())
                p = doc.add_paragraph(style="Caption")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_runs(p, f"Hình {last_h4_number}-{fig_seq}. {desc}")
                continue
            # paragraph thường
            p = doc.add_paragraph()
            add_runs(p, text)
            continue

    doc.save(docx_path)
    return doc, blocks


if __name__ == "__main__":
    # Mặc định: lấy file mẫu trong examples/v1/, xuất .docx ra thư mục hiện tại.
    _default_md = os.path.normpath(
        os.path.join(HERE, "..", "..", "examples", "v1", "SRS_Sample.md"))
    md = sys.argv[1] if len(sys.argv) > 1 else _default_md
    out = sys.argv[2] if len(sys.argv) > 2 else os.path.join(os.getcwd(),
                                                             "SRS_Sample.docx")
    doc, blocks = build(md, out)
    n_h = sum(1 for k, *_ in blocks if k == "heading")
    n_t = sum(1 for k, *_ in blocks if k == "table")
    print(f"Built: {out}")
    print(f"  headings={n_h}, tables={n_t}, total blocks={len(blocks)}")
